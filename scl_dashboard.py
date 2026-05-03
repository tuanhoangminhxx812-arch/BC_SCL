import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import re

# --- Cấu hình trang Streamlit ---
st.set_page_config(page_title="Báo Cáo Kế Toán SCL", layout="wide", initial_sidebar_state="collapsed")

# ============================================================
# CSS RESPONSIVE – Tương thích máy tính & điện thoại
# ============================================================
st.markdown("""
<style>
/* ── Google Font ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* ── Ẩn header mặc định Streamlit ── */
#MainMenu { visibility: hidden; }
footer    { visibility: hidden; }
header    { visibility: hidden; }

/* ── Giảm padding tổng thể trên mobile ── */
.block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 2rem !important;
    padding-left: 1.2rem !important;
    padding-right: 1.2rem !important;
    max-width: 100% !important;
}

/* ── Metric cards ── */
[data-testid="metric-container"] {
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid #334155;
    border-radius: 12px;
    padding: 1rem 1.2rem;
    box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    transition: transform 0.2s;
}
[data-testid="metric-container"]:hover {
    transform: translateY(-2px);
}
[data-testid="metric-container"] label {
    color: #94a3b8 !important;
    font-size: 0.8rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #f1f5f9 !important;
    font-size: 1.3rem !important;
    font-weight: 700 !important;
    word-break: break-word;
}

/* ── Title ── */
h1 { font-size: clamp(1.2rem, 4vw, 2rem) !important; }
h2 { font-size: clamp(1rem, 3vw, 1.5rem) !important; }
h3 { font-size: clamp(0.9rem, 2.5vw, 1.2rem) !important; }

/* ── Buttons ── */
.stButton > button {
    border-radius: 8px !important;
    font-weight: 600 !important;
    transition: all 0.2s !important;
}

/* ── Dataframe responsive ── */
[data-testid="stDataFrame"] {
    overflow-x: auto !important;
    -webkit-overflow-scrolling: touch;
}
[data-testid="stDataFrame"] table {
    font-size: clamp(0.65rem, 1.5vw, 0.9rem) !important;
}

/* ── Divider ── */
hr { border-color: #334155 !important; margin: 1rem 0 !important; }

/* ============================================
   MOBILE – màn hình ≤ 768px
   Streamlit columns sẽ tự xuống dòng
   ============================================ */
@media (max-width: 768px) {

    /* Padding nhỏ lại */
    .block-container {
        padding-left: 0.6rem !important;
        padding-right: 0.6rem !important;
        padding-top: 0.8rem !important;
    }

    /* Các cột Streamlit stack dọc */
    [data-testid="column"] {
        width: 100% !important;
        flex: 0 0 100% !important;
        min-width: 100% !important;
    }

    /* Metric nhỏ hơn */
    [data-testid="metric-container"] {
        padding: 0.7rem 0.9rem;
        margin-bottom: 0.5rem;
    }
    [data-testid="metric-container"] [data-testid="stMetricValue"] {
        font-size: 1.1rem !important;
    }

    /* Chart full width */
    .stpyplot, iframe {
        width: 100% !important;
    }

    /* Title nhỏ hơn */
    h1 { font-size: 1.1rem !important; line-height: 1.4; }
    h2 { font-size: 1rem !important; }
    h3 { font-size: 0.9rem !important; }

    /* Nút full width trên mobile */
    .stButton > button, .stDownloadButton > button {
        width: 100% !important;
        font-size: 0.85rem !important;
        padding: 0.5rem !important;
    }

    /* Bảng cuộn ngang */
    [data-testid="stDataFrame"] {
        font-size: 0.65rem !important;
    }

    /* Alert / info box */
    .stAlert { font-size: 0.8rem !important; }

    /* Expander */
    .streamlit-expanderHeader { font-size: 0.85rem !important; }
}

/* ============================================
   TABLET – màn hình 769px – 1024px
   ============================================ */
@media (min-width: 769px) and (max-width: 1024px) {
    [data-testid="metric-container"] [data-testid="stMetricValue"] {
        font-size: 1.1rem !important;
    }
    h1 { font-size: 1.4rem !important; }
}
</style>
""", unsafe_allow_html=True)


# ============================================================
# HÀM ĐỌC FILE PM_092 – TỔNG SỐ DƯ CUỐI KỲ THEO CÔNG TRÌNH
# ============================================================
@st.cache_data
def load_pm092():
    """
    Đọc file PM_092_*.xlsx trong cùng thư mục, parse và trả về
    dict { mã_công_trình: tổng_số_dư_cuối_kỳ } từ dòng
    'Tổng số dư cuối kỳ _ CÔNG TRÌNH' tương ứng với từng công trình.
    """
    base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()

    # Tìm tất cả file PM_092*.xlsx – lấy file mới nhất (theo thời gian sửa)
    pm_files = [
        os.path.join(base_dir, f)
        for f in os.listdir(base_dir)
        if f.upper().startswith("PM_092") and f.lower().endswith(".xlsx")
    ]

    if not pm_files:
        return {}, None  # không có file

    pm_path = max(pm_files, key=os.path.getmtime)
    pm_filename = os.path.basename(pm_path)

    raw = pd.read_excel(pm_path, sheet_name=0, header=None)

    result = {}
    current_ma_ct = None

    for _, row in raw.iterrows():
        cell0 = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ""

        # Nhận diện dòng "Công trình: MACT - Tên..."
        if cell0.startswith("Công trình:"):
            # Trích mã công trình (phần đầu sau "Công trình: ", trước " - ")
            parts = cell0.replace("Công trình:", "").strip()
            ma = parts.split(" - ")[0].strip()
            current_ma_ct = ma

        # Nhận diện dòng "Tổng số dư cuối kỳ _ CÔNG TRÌNH"
        elif "Tổng số dư cuối kỳ _ CÔNG TRÌNH" in cell0 and current_ma_ct:
            # Giá trị nằm ở cột 4 (index 4)
            val = row.iloc[4]
            if pd.notna(val):
                try:
                    result[current_ma_ct] = float(val)
                except (ValueError, TypeError):
                    pass
            current_ma_ct = None  # reset – tránh gán nhầm công trình tiếp theo

    return result, pm_filename


# --- Đọc dữ liệu chính ---
@st.cache_data
def load_data():
    base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    file_path = os.path.join(base_dir, "Tong Hop.xlsx")

    if not os.path.exists(file_path):
        local_path = r"D:\HOC A.I\KT SCL\BC SCL\Tong Hop.xlsx"
        if os.path.exists(local_path):
            file_path = local_path

    if not os.path.exists(file_path):
        st.error("Không tìm thấy file dữ liệu 'Tong Hop.xlsx'.")
        return pd.DataFrame()

    df = pd.read_excel(file_path, sheet_name="Sheet1")
    df = df[df['Mã công trình'].notna()]
    df['Giá trị khái toán'] = df['Giá trị khái toán'].fillna(0)
    df['Giá trị thực hiện'] = df['Giá trị thực hiện'].fillna(0)
    df['Giá trị quyết toán'] = df['Giá trị quyết toán'].fillna(0)
    return df


# ============================================================
# LOAD DỮ LIỆU
# ============================================================
df = load_data()
pm_dict, pm_filename = load_pm092()

if df.empty:
    st.stop()

# --- Áp dụng Tổng số dư cuối kỳ _ CÔNG TRÌNH vào Giá trị thực hiện ---
so_cong_trinh_cap_nhat = 0
if pm_dict:
    def update_thuc_hien(row):
        ma = str(row['Mã công trình']).strip()
        if ma in pm_dict:
            return pm_dict[ma]
        return row['Giá trị thực hiện']

    df_updated = df.copy()
    df_updated['Giá trị thực hiện'] = df_updated.apply(update_thuc_hien, axis=1)
    so_cong_trinh_cap_nhat = sum(
        1 for ma in df['Mã công trình'].astype(str).str.strip() if ma in pm_dict
    )
    df = df_updated

# --- Tính toán các chỉ số ---
tong_khai_toan = df['Giá trị khái toán'].sum()
tong_thuc_hien = df['Giá trị thực hiện'].sum()
tong_quyet_toan = df['Giá trị quyết toán'].sum()
ty_le_giai_ngan = (tong_thuc_hien / tong_khai_toan * 100) if tong_khai_toan > 0 else 0

# ============================================================
# GIAO DIỆN DASHBOARD
# ============================================================
st.title("📊 BÁO CÁO TỔNG HỢP & PHÂN TÍCH QUẢN TRỊ CHI PHÍ SCL")

# --- Thanh nút điều khiển ---
col_title, col_btn1, col_btn2 = st.columns([6, 2, 2])
with col_btn1:
    if st.button("🔄 Tải lại Tong Hop", type="secondary", use_container_width=True):
        load_data.clear()
        st.rerun()
with col_btn2:
    if st.button("📥 Cập nhật từ PM_092", type="primary", use_container_width=True):
        load_pm092.clear()
        load_data.clear()
        st.rerun()

# --- Banner thông tin file PM_092 ---
if pm_filename:
    st.success(
        f"✅ **Nguồn dữ liệu thực hiện:** `{pm_filename}` – "
        f"Đã cập nhật **{so_cong_trinh_cap_nhat}/{len(df)}** công trình từ "
        f"cột 'Tổng số dư cuối kỳ _ CÔNG TRÌNH'"
    )
else:
    st.warning("⚠️ Chưa tìm thấy file **PM_092\\*.xlsx** trong thư mục. "
               "Giá trị thực hiện đang dùng dữ liệu từ Tong Hop.xlsx.")

# --- Hiển thị bảng mapping PM_092 ---
if pm_dict:
    with st.expander("🔍 Chi tiết số dư cuối kỳ theo công trình từ PM_092 (click để xem)"):
        pm_df = pd.DataFrame([
            {"Mã công trình": ma, "Tổng số dư cuối kỳ _ CÔNG TRÌNH (đ)": f"{val:,.0f}"}
            for ma, val in pm_dict.items()
        ])
        st.dataframe(pm_df, use_container_width=True)

st.markdown("---")

# --- Các chỉ số tổng quan ---
col1, col2, col3, col4 = st.columns(4)
col1.metric("Tổng Số Công Trình", len(df))
col2.metric("Tổng Giá Trị Khái Toán", f"{tong_khai_toan:,.0f} đ")
col3.metric("Tổng Giá Trị Thực Hiện", f"{tong_thuc_hien:,.0f} đ")
col4.metric("Tỷ Lệ Giải Ngân", f"{ty_le_giai_ngan:.2f} %")

st.markdown("---")
st.subheader("📈 Sơ đồ trực quan hóa dữ liệu")

col_chart1, col_chart2 = st.columns(2)

with col_chart1:
    st.markdown("**1. Tỷ trọng trạng thái dự án**")
    status_counts = df['Trạng thái'].value_counts()

    fig1, ax1 = plt.subplots(figsize=(7, 4))
    colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99']
    ax1.pie(status_counts, labels=status_counts.index, autopct='%1.1f%%',
            startangle=90, colors=colors[:len(status_counts)],
            wedgeprops={'edgecolor': 'white'})
    ax1.axis('equal')
    st.pyplot(fig1)

with col_chart2:
    st.markdown("**2. Top dự án có mức ngân sách cao nhất (Khái toán vs Thực hiện)**")
    df_sorted = df.sort_values(by='Giá trị khái toán', ascending=False).head(5)

    fig2, ax2 = plt.subplots(figsize=(8, 4))
    x = range(len(df_sorted))
    width = 0.35

    khai_toan_ty = df_sorted['Giá trị khái toán'] / 1e9
    thuc_hien_ty = df_sorted['Giá trị thực hiện'] / 1e9

    ax2.bar([i - width/2 for i in x], khai_toan_ty, width, label='Khái toán (Tỷ đ)', color='#2171b5')
    ax2.bar([i + width/2 for i in x], thuc_hien_ty, width, label='Thực hiện (Tỷ đ)', color='#fd8d3c')

    ax2.set_xticks(x)
    m_cong_trinh = df_sorted['Mã công trình'].tolist()
    ax2.set_xticklabels(m_cong_trinh, rotation=30, ha="right")
    ax2.legend()
    ax2.grid(axis='y', linestyle='--', alpha=0.7)

    st.pyplot(fig2)

st.markdown("---")
st.subheader("📋 Bảng số liệu chi tiết các dự án SCL")

# Chuẩn bị dữ liệu
df_display = df[['Mã công trình', 'Tên công trình', 'Trạng thái',
                  'Giá trị khái toán', 'Giá trị thực hiện', 'Giá trị quyết toán']].copy()
df_display['Nguồn TH'] = df_display['Mã công trình'].apply(
    lambda ma: '📥 PM_092' if str(ma).strip() in pm_dict else '📄 Tong Hop'
)
for col in ['Giá trị khái toán', 'Giá trị thực hiện', 'Giá trị quyết toán']:
    df_display[col] = df_display[col].apply(lambda x: f"{x:,.0f}")

# ── Màu badge trạng thái ──
STATUS_COLOR = {
    'Đang thi công':          ('#22c55e', '#052e16'),
    'Lập kế hoạch đầu thầu': ('#3b82f6', '#0c1a3a'),
    'Lập PAKT-Tổng dự toán':  ('#f59e0b', '#2d1a00'),
    'Hoàn thành':             ('#a855f7', '#1a0a2e'),
    'Quyết toán':             ('#06b6d4', '#042329'),
}
def status_badge(trang_thai):
    color, bg = STATUS_COLOR.get(str(trang_thai), ('#94a3b8', '#1e293b'))
    return (f'<span style="background:{bg};color:{color};border:1px solid {color};'
            f'border-radius:20px;padding:2px 10px;font-size:0.75rem;'
            f'font-weight:600;white-space:nowrap;">{trang_thai}</span>')

# ── Build HTML table (Desktop) + Cards (Mobile) ──
table_rows = ""
cards_html = ""

for _, row in df_display.iterrows():
    ma        = row['Mã công trình']
    ten       = row['Tên công trình']
    tt        = row['Trạng thái']
    kt        = row['Giá trị khái toán']
    th        = row['Giá trị thực hiện']
    qt        = row['Giá trị quyết toán']
    nguon     = row['Nguồn TH']
    badge     = status_badge(tt)

    # Desktop row
    table_rows += f"""
<tr>
  <td style="font-weight:600;color:#60a5fa;">{ma}</td>
  <td style="max-width:220px;">{ten}</td>
  <td>{badge}</td>
  <td style="text-align:right;font-variant-numeric:tabular-nums;">{kt}</td>
  <td style="text-align:right;font-variant-numeric:tabular-nums;color:#34d399;">{th}</td>
  <td style="text-align:right;font-variant-numeric:tabular-nums;color:#a78bfa;">{qt}</td>
  <td style="text-align:center;">{nguon}</td>
</tr>
"""

    # Mobile card
    cards_html += f"""
<div class="mobile-card">
  <div class="mc-header">
    <span class="mc-ma">{ma}</span>
    {badge}
  </div>
  <div class="mc-ten">{ten}</div>
  <div class="mc-row"><span class="mc-label">Khái toán</span>
    <span class="mc-val">{kt} đ</span></div>
  <div class="mc-row"><span class="mc-label">Thực hiện</span>
    <span class="mc-val mc-green">{th} đ</span></div>
  <div class="mc-row"><span class="mc-label">Quyết toán</span>
    <span class="mc-val mc-purple">{qt} đ</span></div>
  <div class="mc-row"><span class="mc-label">Nguồn TH</span>
    <span class="mc-val">{nguon}</span></div>
</div>
"""

html_table = f"""
<style>
/* ── DESKTOP TABLE ── */
.scl-wrap {{
  overflow-x: auto;
  -webkit-overflow-scrolling: touch;
  border-radius: 12px;
}}
.scl-table {{
  width: 100%;
  border-collapse: collapse;
  font-size: 0.88rem;
  background: #0f172a;
  color: #e2e8f0;
}}
.scl-table thead tr {{
  background: #1e293b;
}}
.scl-table th {{
  padding: 10px 12px;
  text-align: left;
  font-size: 0.75rem;
  font-weight: 700;
  letter-spacing: 0.06em;
  text-transform: uppercase;
  color: #94a3b8;
  white-space: nowrap;
  border-bottom: 2px solid #334155;
}}
.scl-table td {{
  padding: 9px 12px;
  border-bottom: 1px solid #1e293b;
  vertical-align: middle;
  line-height: 1.4;
}}
.scl-table tbody tr:hover td {{
  background: #1e293b;
}}

/* ── MOBILE CARDS (ẩn trên desktop, hiện trên mobile) ── */
.mobile-cards  {{ display: none; }}
.desktop-table {{ display: block; }}

.mobile-card {{
  background: #1e293b;
  border: 1px solid #334155;
  border-radius: 14px;
  padding: 14px 16px;
  margin-bottom: 12px;
}}
.mc-header {{
  display: flex;
  align-items: center;
  gap: 10px;
  margin-bottom: 8px;
}}
.mc-ma {{
  font-size: 0.95rem;
  font-weight: 700;
  color: #60a5fa;
}}
.mc-ten {{
  font-size: 0.85rem;
  color: #cbd5e1;
  margin-bottom: 10px;
  line-height: 1.4;
}}
.mc-row {{
  display: flex;
  justify-content: space-between;
  align-items: center;
  padding: 5px 0;
  border-bottom: 1px solid #334155;
}}
.mc-row:last-child {{ border-bottom: none; }}
.mc-label {{
  font-size: 0.75rem;
  color: #94a3b8;
  font-weight: 500;
  text-transform: uppercase;
  letter-spacing: 0.05em;
}}
.mc-val {{
  font-size: 0.85rem;
  font-weight: 600;
  color: #e2e8f0;
  text-align: right;
  font-variant-numeric: tabular-nums;
}}
.mc-green {{ color: #34d399 !important; }}
.mc-purple {{ color: #a78bfa !important; }}

@media (max-width: 768px) {{
  .mobile-cards  {{ display: block; }}
  .desktop-table {{ display: none; }}
}}
</style>

<!-- Desktop table -->
<div class="desktop-table scl-wrap">
  <table class="scl-table">
    <thead>
      <tr>
        <th>Mã CT</th>
        <th>Tên công trình</th>
        <th>Trạng thái</th>
        <th style="text-align:right">Khái toán (đ)</th>
        <th style="text-align:right">Thực hiện (đ)</th>
        <th style="text-align:right">Quyết toán (đ)</th>
        <th style="text-align:center">Nguồn TH</th>
      </tr>
    </thead>
    <tbody>
{table_rows}
    </tbody>
  </table>
</div>

<!-- Mobile cards -->
<div class="mobile-cards">
{cards_html}
</div>
"""

st.markdown(html_table, unsafe_allow_html=True)


# --- Phân tích của Kế toán trưởng ---
st.markdown("---")
st.subheader("⚠️ Phân tích rủi ro & Đề xuất (Trình Ban Giám Đốc)")

if ty_le_giai_ngan < 30:
    nhan_xet_giai_ngan = "Ở mức **báo động đỏ** (Trễ tiến độ giải ngân)"
    kl_giai_ngan = "Sự chênh lệch lớn giữa Ngân sách và Thực tế cho thấy các thủ tục chuẩn bị hồ sơ thanh toán đang bị đình trệ nghiêm trọng."
elif ty_le_giai_ngan < 70:
    nhan_xet_giai_ngan = "Ở mức **trung bình** (Cần đẩy nhanh hơn)"
    kl_giai_ngan = "Tiến độ giải ngân đang được thực hiện nhưng cần đốc thúc thêm để rải đều trong năm, hoàn thành đúng mục tiêu dòng tiền."
else:
    nhan_xet_giai_ngan = "Ở mức **rất tốt** (Hoàn thành theo bám sát kế hoạch)"
    kl_giai_ngan = "Các công tác thi công và nghiệm thu hồ sơ đang phối hợp rất nhịp nhàng, đảm bảo tính pháp lý và giảm tải rủi ro dồn khối lượng vào cuối năm."

so_du_an_0 = len(df[df['Giá trị thực hiện'] == 0])
so_du_an_quyet_toan = len(df[df['Giá trị quyết toán'] > 0])
tong_du_an = len(df)

if so_du_an_0 > 0:
    nhan_xet_0 = f"Báo cáo cho thấy có **{so_du_an_0}** dự án hoàn toàn chưa ghi nhận chứng từ chi phí dở dang ('Giá trị thực hiện' = 0đ)."
    kl_chung_tu = "Cần rà soát chéo lượng công trình này ngay. Xem đây là do thực sự chưa triển khai ngoài hiện trường, hay kỹ thuật đã cho làm nhưng nhà thầu chây ỳ chưa lập hồ sơ nghiệm thu. Tránh tình trạng nợ đọng, thi công xong mà sổ sách không có chứng từ."
else:
    nhan_xet_0 = f"Rất tốt, 100% ({tong_du_an}/{tong_du_an}) dự án đều đã có hồ sơ ghi nhận Khối lượng thực hiện ban đầu."
    kl_chung_tu = "Sự phối hợp cập nhật chứng từ giữa phòng Kỹ thuật và Kế toán đang bám sát thực tế, không có tình trạng bị trễ nhịp hay bỏ quên dự án."

if so_du_an_quyet_toan == 0:
    nhan_xet_qt = "Đồng thời, chưa có dự án nào chuyển sang bước 'Giá trị quyết toán'."
elif so_du_an_quyet_toan == tong_du_an:
    nhan_xet_qt = "Tuyệt vời, tất cả các dự án đều đã có số liệu Quyết Toán! Quá trình khép sổ tài chính SCL gần như đã trọn vẹn."
else:
    nhan_xet_qt = f"Tiến độ quyết toán: Đã có **{so_du_an_quyet_toan}/{tong_du_an}** dự án có số liệu Quyết toán thành công."

if so_du_an_quyet_toan == tong_du_an:
    kien_nghi = "- Hồ sơ tài chính đã đạt mức độ hoàn thiện cao. Đề nghị các phòng ban chuẩn bị đóng luồng hồ sơ cuối năm và báo cáo Giám đốc."
else:
    kien_nghi = "- Đẩy nhanh tiến độ hoàn công chuyển các công trình thành Quyết Toán (QT).\n- Liên tục tổ chức đối chiếu công nợ khối lượng dở dang hàng tháng giữa kế toán và kỹ thuật."

analysis_text = "Dưới đây là phần trình bày tổng hợp các chỉ số đánh giá chuyên môn về mặt quản trị tài chính doanh nghiệp:\n\n"

analysis_text += f"""**1. Tỷ lệ giải ngân: {nhan_xet_giai_ngan}**
- Tổng quy mô vốn khái toán cho {tong_du_an} công trình là hơn **{tong_khai_toan/1e9:,.1f} tỷ đồng**.
- Tuy nhiên, giá trị khối lượng thực hiện ghi nhận trên sổ sách là **{tong_thuc_hien/1e9:,.2f} tỷ đồng**, đạt mức **{ty_le_giai_ngan:.2f}%**.
=> Kết luận: {kl_giai_ngan}

**2. Công tác đồng bộ giữa thực địa và chứng từ (Phòng Kỹ thuật vs Kế toán)**
- {nhan_xet_0}
- {nhan_xet_qt}
=> Kết luận: {kl_chung_tu}

**🔴 KIẾN NGHỊ TỪ KẾ TOÁN TRƯỞNG:**
{kien_nghi}
"""

st.markdown(analysis_text)

# --- Hàm tạo báo cáo Word ---
def export_word_report():
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_paragraph('CÔNG TY ĐIỆN LỰC VŨNG TÀU\nPHÒNG TÀI CHÍNH KẾ TOÁN')

    title = doc.add_paragraph('\nBÁO CÁO PHÂN TÍCH TÌNH HÌNH THỰC HIỆN KẾ HOẠCH TÀI CHÍNH\nCÔNG TÁC SỬA CHỮA LỚN NĂM 2026\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)

    p = doc.add_paragraph()
    run_kinh_gui = p.add_run('Kính gửi: ')
    run_kinh_gui.bold = True
    p.add_run('Ông (Bà) Giám đốc Công ty')

    # Ghi nguồn dữ liệu vào Word
    src_note = f"(Nguồn dữ liệu thực hiện: {pm_filename})" if pm_filename else "(Nguồn dữ liệu thực hiện: Tong Hop.xlsx)"
    doc.add_paragraph(
        f"Căn cứ vào dữ liệu tổng hợp về tình hình thực hiện kế hoạch các dự án sửa chữa lớn {src_note}, "
        "trên cương vị Kế toán trưởng, tôi xin báo cáo các số liệu tài chính quan trọng và các điểm bất ổn "
        "cần Giám đốc khẩn trương chỉ đạo như sau:"
    )

    doc.add_paragraph('I. BẢNG TỔNG HỢP SỐ LIỆU TÀI CHÍNH:', style='Heading 3')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ tiêu'
    hdr_cells[1].text = 'Kế hoạch/Khái toán (VNĐ)'
    hdr_cells[2].text = 'Giá trị thực hiện (VNĐ)'
    hdr_cells[3].text = 'Tỉ lệ hoàn thành (%)'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = 'Toàn bộ công trình SCL'
    row_cells[1].text = f"{tong_khai_toan:,.0f}"
    row_cells[2].text = f"{tong_thuc_hien:,.0f}"
    row_cells[3].text = f"{ty_le_giai_ngan:.2f}%"

    doc.add_paragraph()

    doc.add_paragraph('II. PHÂN TÍCH ĐÁNH GIÁ & CẢNH BÁO BẤT ỔN:', style='Heading 3')

    import re as _re
    parts = analysis_text.split('\n')
    for p_text in parts:
        clean = p_text.strip()
        if not clean:
            continue
        # Bỏ ký hiệu Markdown trước khi ghi vào Word
        plain = _re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
        is_bold = (clean.startswith("🔴") or "CẢNH BÁO" in clean
                   or clean.startswith("**") or clean.startswith("=>"))
        p_docx = doc.add_paragraph(plain)
        if is_bold:
            run_bold = p_docx.runs[0] if p_docx.runs else p_docx.add_run(plain)
            run_bold.bold = True

    doc.add_paragraph()

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = p_sig.add_run('Vũng Tàu, ngày ...... tháng ...... năm ...... \n')
    run_date.italic = True
    run_title = p_sig.add_run('KẾ TOÁN TRƯỞNG\n\n\n\n\n')
    run_title.bold = True
    p_sig.add_run('(Đã ký)')

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


st.markdown("### 📥 Tải Xuất báo cáo chính thức")

# Thư mục lưu file — cùng thư mục với script
_base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()

# Tạo sẵn file Word vào session_state để tránh lỗi khi Streamlit re-run lúc bấm nút
if 'word_report_bytes' not in st.session_state:
    st.session_state['word_report_bytes'] = None
if 'word_saved_path' not in st.session_state:
    st.session_state['word_saved_path'] = None

col_dl1, col_dl2 = st.columns([3, 1])
with col_dl1:
    if st.button("⚙️ Tạo / Cập nhật File Báo Cáo Word", use_container_width=True):
        with st.spinner("Đang tạo file Word..."):
            try:
                from datetime import datetime
                word_bytes = export_word_report()
                st.session_state['word_report_bytes'] = word_bytes

                # Lưu thẳng vào thư mục chứa file báo cáo
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                save_filename = f"Bao_Cao_SCL_KeToanTruong_{timestamp}.docx"
                save_path = os.path.join(_base_dir, save_filename)
                with open(save_path, "wb") as f:
                    f.write(word_bytes)
                st.session_state['word_saved_path'] = save_path

                st.success(f"✅ Đã tạo & lưu file vào:\n\n📁 `{save_path}`")
            except Exception as e:
                st.error(f"❌ Lỗi tạo file Word: {e}")

    # Hiển thị đường dẫn file đã lưu (nếu có)
    if st.session_state.get('word_saved_path') and os.path.exists(st.session_state['word_saved_path']):
        saved = st.session_state['word_saved_path']
        st.info(f"📂 File đã lưu tại: `{saved}`")

with col_dl2:
    if st.session_state.get('word_report_bytes'):
        st.download_button(
            label="📄 Tải Xuống (.docx)",
            data=st.session_state['word_report_bytes'],
            file_name="Bao_Cao_SCL_KeToanTruong.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("Bấm '⚙️ Tạo File' trước")

# Chạy ứng dụng bằng lệnh: streamlit run "D:\HOC A.I\KT SCL\BC SCL\scl_dashboard.py"
